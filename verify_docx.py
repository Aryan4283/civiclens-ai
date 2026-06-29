from docx import Document

try:
    doc = Document('CivicLens_Submission_Report.docx')
    paragraphs = len(doc.paragraphs)
    
    # In python-docx, images are stored in inline shapes
    images = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if r._element.xpath('.//pic:pic'):
                images += len(r._element.xpath('.//pic:pic'))
    
    print(f"Verification Success!")
    print(f"Total Paragraphs/Headings: {paragraphs}")
    print(f"Total Images successfully embedded: {images}")
    
except Exception as e:
    print(f"Error reading docx: {e}")

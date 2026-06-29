import os
import re
from docx import Document
from docx.shared import Inches

def create_doc():
    doc = Document()
    
    with open('submission_report.md', 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if it's a heading
        if line.startswith('# '):
            doc.add_heading(line[2:].replace('**', ''), 0)
        elif line.startswith('## '):
            doc.add_heading(line[3:].replace('**', ''), 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].replace('**', ''), 2)
            
        # Check if it's an image
        elif line.startswith('![') and '](' in line:
            # Extract image path
            match = re.search(r'!\[.*?\]\((.*?)\)', line)
            if match:
                img_path = match.group(1)
                if img_path.startswith('./'):
                    img_path = img_path[2:]
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(5.5))
                    except Exception as e:
                        doc.add_paragraph(f"Error loading image: {img_path}")
                else:
                    doc.add_paragraph(f"[Image not found: {img_path}]")
        
        # Check if it's a list item
        elif line.startswith('- ') or line.startswith('* '):
            clean_line = line[2:].replace('**', '')
            doc.add_paragraph(clean_line, style='List Bullet')
            
        else:
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)
            
    doc.save('CivicLens_Submission_Report.docx')
    print("Successfully generated CivicLens_Submission_Report.docx!")

if __name__ == "__main__":
    create_doc()
